from PyQt5 import QtCore, QtWidgets
from PyQt5.QtCore import QEvent, QObject
from PyQt5.QtWidgets import (
    QMenu,
    QTableWidgetItem,
    QDialogButtonBox,
    QHeaderView,
    QLabel,
    QTableWidget,
    QWidget,
    QGridLayout,
)

import docx

# Load ui
from ui.ui_main import Ui_MainWindow

# Database
from database.Customer import Customer
from database.Product import Product
from database.Invoice import Invoice

# Load View
from views.AddCusView import AddCusView
from views.AddProductView import AddProductView
from views.AddProduct2InvoiceView import AddProduct2Invoice
from views.ErrorView import WarningView
from views.TableWidget import InvoiceTable




class MainWindow(QtWidgets.QMainWindow):
    def __init__(self, *args, **kwargs):
        super(MainWindow, self).__init__(*args, **kwargs)

        # Load the UI Page
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        """
        CUSTOMER
        """
        # Load all cus data
        cus = Customer()
        datas = cus.get_all()
        for d in datas:
            self.ui.cus_list.addItems(d)

        # search user bar
        self.ui.cus_search_bar.textChanged.connect(self.on_cus_search_bar)
        self.ui.cus_list.itemDoubleClicked.connect(self.on_cus_list_widget)
        self.ui.cus_list.installEventFilter(self)

        # Add cus btn
        self.ui.add_cus.clicked.connect(self.on_add_cus)

        """
        PRODUCT
        """
        # Load product data
        product = Product()
        products = product.get_all_p_names()
        for p in products:
            self.ui.p_list.addItems(p)

        # search product bar
        self.ui.p_search_bar.textChanged.connect(self.on_p_search_bar)
        self.ui.p_list.itemDoubleClicked.connect(self.on_p_list_widget)
        self.ui.p_list.installEventFilter(self)

        # Add product btn
        self.ui.add_p.clicked.connect(self.on_add_p)

        """INVOICE
        """
        self.cus_choosed = False
        self.ui.add_p2invoice.clicked.connect(self.on_add_p2invoice)
        self.ui.export2pdf.clicked.connect(self.on_export2pdf)

    def on_export2pdf(self):
        # Create an instance of a word document
        doc = docx.Document()
        # Add a Title to the document
        doc.add_heading(f"Công nợ sổ của {self.cus_name}", 0)

        # Table data in a form of list
        labels = [
            "Tên sản phẩm",
            "Đơn giá (VND)",
            "Số lượng",
            "Tổng (Đơn giá * Số lượng) VND",
        ]

        # Adding data from the list to the table
        for id, create_date in Invoice().getAllInvoice(self.cus_name):
            [(total_p, months)] = Invoice().getTotalFromInvoice(id)
            para = doc.add_paragraph()
            bold_para = para.add_run(
                f"Hóa đơn ngày: {create_date}\tSố tháng thiếu: {months}"
            )
            # bold_para = para.add_run(f"Ngày ngày trả: {create_date}")
            bold_para.bold = True
            # Creating a table object
            table = doc.add_table(rows=1, cols=len(labels))
            table.style = "TableGrid"
            # Adding heading in the 1st row of the table
            row = table.rows[0].cells
            for i in range(len(labels)):
                row[i].text = labels[i]

            for i_data in Invoice().getAllProductFromInvoice(id):
                # Adding a row and then adding data in it.
                row = table.add_row().cells
                # # Converting id to string as table can only take string input
                for i in range(len(i_data)):
                    row[i].text = str(i_data[i])

            doc.add_paragraph(
                f"Tổng tiền trong ngày: {Invoice().seperated_by.format(total_p)} VND \t Tiền lãi: {Invoice().seperated_by.format(float(total_p) * float(months)* Invoice().interest_rate)} VND"
            )
        # Now save the document to a location
        doc.save(
            "D:\\banhb\Documents\Projects\\2023\simple-store\docx template\Invoice.docx"
        )
        print("EXPORT SUCCESFULLY")

    def on_add_p2invoice(self):
        # Create a buffer for query
        query = dict()

        if self.cus_choosed:
            dlg = AddProduct2Invoice()
            if dlg.exec():
                query["c_name"] = self.cus_name
                query["p_name"] = dlg.ui.p_name_input.text()
                query["p_quantity"] = int(dlg.ui.p_quantity_input.text())
                query["i_create_date"] = dlg.ui.create_date.text()

                # Check if new price
                if len(dlg.ui.new_unit_price_input.text()):
                    query["p_unit_price"] = float(
                        dlg.p.de_seperated(dlg.ui.new_unit_price_input.text())
                    )
                else:
                    query["p_unit_price"] = float(
                        dlg.p.de_seperated(dlg.ui.unit_price_input.text())
                    )

                dlg.addp2Invoice(query)

                # update Invoice
                self.update_invoices_table_widget(query["c_name"])
                # Update total
                self.update_total(query["c_name"])

            else:
                print("Cancel")
        else:
            WarningView(
                "Vui lòng chọn khách hàng!!!\nNhấp chuột 2 lần vào bảng danh sách khách hàng."
            )

    # https://www.youtube.com/watch?v=2Ie-EBwJOZU
    def eventFilter(self, source: QObject, event: QEvent) -> bool:
        if event.type() == QEvent.ContextMenu and source is self.ui.cus_list:
            cus_menu = QMenu()
            cus_menu.addAction("Edit", self.cus_edit)
            cus_menu.addAction("Delete", self.cus_delete)
            cus_menu.exec_(event.globalPos())
            return True
        elif event.type() == QEvent.ContextMenu and source is self.ui.p_list:
            p_menu = QMenu()
            p_menu.addAction("Edit", self.p_edit)
            p_menu.addAction("Delete", self.p_delete)
            p_menu.exec_(event.globalPos())
            return True

        return super().eventFilter(source, event)

    """
    CRUD PRODUCT
    """

    def on_add_p(self):
        print("add_p clicked")
        dlg = AddProductView()
        dlg.ui.buttonBox.button(QDialogButtonBox.Ok).setEnabled(False)
        p_data = dict()

        if dlg.exec():
            p_data["name"] = dlg.ui.p_name_input.text()
            p_data["unit_price"] = dlg.ui.unit_price_input.text()
            if len(p_data["unit_price"]) > 0 and len(p_data["name"]) > 0:
                dlg.add(p_data)
                print(f"Added {p_data}")
            else:
                print("Input Nothing!!")
        else:
            print("Cancel")

    def p_edit(self):
        pass

    def p_delete(self):
        pass

    """
    CRUD CUSTOMER
    """

    def on_add_cus(self):
        print("add_cus clicked")
        dlg = AddCusView()
        cus_data = dict()

        if dlg.exec():
            cus_data["name"] = dlg.ui.cus_name_input.text()
            if len(cus_data["name"]) > 0:
                print(type(cus_data["name"]))
                dlg.add(cus_data)
                print(f"Added customer: {cus_data}")
            else:
                print("Input Nothing!!")
        else:
            print("Cancel")

    def cus_delete(self):
        WarningView(
            "Xóa khách hàng sẽ xóa hết tất cả hóa đơn liên quan\nBạn có chắc chắn không?"
        )

    def cus_edit(self):
        print("cus_edit")

    def on_p_list_widget(self, text):
        print(f"Product '{text.text()}' is selected")
        p = Product()
        p_infos = p.get_per_product(text.text())
        for p_info in p_infos:
            self.ui.p_name.setText(f"# {p_info[0]}")
            # self.ui.p_unit_price.setText(f"**{p.seperated_by.format(p_info[1])}**")
            self.ui.p_unit_price.setText(f"**{(p_info[1])}**")
            self.ui.p_location.setText(f"**{p_info[2]}**")

    def on_p_search_bar(self, text):
        self.ui.p_list.clear()
        products = Product().get_p_name(value=text)
        for p in products:
            self.ui.p_list.addItems(p)

    def on_cus_list_widget(self, text):
        # Add Condition for p2Invoice
        self.cus_choosed = True
        self.cus_name = text.text()  # Use this to query p2Invoice

        # UI setup
        print(f"Customer '{text.text()}' is selected")
        self.ui.invoice_name.setText(f"# Các hóa đơn thiếu của {text.text()}")

        invoices = Invoice().getAllInvoice(self.cus_name)

        widget = QWidget()
        self.grid = QGridLayout()

        # Manage multiples layout with this
        self.table_list = []
        self.add_btn_list = []
        self.update_btn_list = []
        self.remove_btn_list = []
        self.handleCellChanged_list = []
        # Create view
        for invoice_id, invoice_create_date in invoices:
            self.createTemplate(invoice_id, invoice_create_date)

        # table_list
        for i in range(len(self.table_list)):
            # self.table_list[i].cellChanged.connect(self.handleCellChanged)
            self.table_list[i].cellClicked.connect(self.handleCellChanged)

        widget.setLayout(self.grid)
        self.ui.scroll_infos.setWidget(widget)

    def createTemplate(self, i_id, i_create_date):
        # Table setup
        invoice_table = InvoiceTable(i_id)        
        invoice_table.setMinimumSize(QtCore.QSize(0, 200))
        invoice_table.setSortingEnabled(True)

        create_date = QLabel(f"**Hóa đơn ngày: {f'{i_create_date}'}**")
        create_date.setTextFormat(QtCore.Qt.MarkdownText)

        self.grid.addWidget(create_date)
        self.grid.addWidget(invoice_table)

        # Insert data into this template
        labels = [
            "Tên sản phẩm",
            "Đơn giá (VND)",
            "Số lượng",
            "Tổng (Đơn giá * Số lượng) VND",
        ]
        self.update_table(
            invoice_table, Invoice().getAllProductFromInvoice(i_id), labels
        )

    def update_total(self, name):
        [(total_p, total_interest)] = Invoice().getTotal(name)

        # if Customer did bought anything.
        if total_interest is None and total_p is None:
            [(total_p, total_interest)] = [(0, 0)]

    def update_table(self, table, invoice_data, labels):
        table.setRowCount(0)
        table.setColumnCount(len(labels))
        table.setHorizontalHeaderLabels(labels)

        # table.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeToContents)

        for i in range(len(labels)):
            table.horizontalHeader().setSectionResizeMode(i, QHeaderView.Stretch)

        # DONT KNOW HOW, BUT IT WORKS https://www.youtube.com/watch?v=_QKVDfVyRbM
        for row_number, row_data in enumerate(invoice_data):
            table.insertRow(row_number)

            for column_number, data in enumerate(row_data):
                table.setItem(row_number, column_number, QTableWidgetItem(str(data)))

    def on_cus_search_bar(self, text):
        self.ui.cus_list.clear()
        for data in Customer().get_name(value=text):
            self.ui.cus_list.addItems(data)
